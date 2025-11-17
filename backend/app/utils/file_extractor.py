"""
File extraction utilities for PDF and DOCX files
"""
from typing import Optional
from PyPDF2 import PdfReader
from docx import Document
import io


class FileExtractor:
    """Extract text content from various file formats"""

    @staticmethod
    async def extract_from_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF file

        Args:
            file_content: PDF file content as bytes

        Returns:
            Extracted text content
        """
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PdfReader(pdf_file)

            text_content = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)

            return "\n\n".join(text_content)
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")

    @staticmethod
    async def extract_from_docx(file_content: bytes) -> str:
        """
        Extract text from DOCX file

        Args:
            file_content: DOCX file content as bytes

        Returns:
            Extracted text content
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)

            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)

            return "\n\n".join(text_content)
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

    @staticmethod
    async def extract_text(file_content: bytes, content_type: str) -> str:
        """
        Extract text from file based on content type

        Args:
            file_content: File content as bytes
            content_type: MIME type of the file

        Returns:
            Extracted text content
        """
        if content_type == "application/pdf":
            return await FileExtractor.extract_from_pdf(file_content)
        elif content_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ]:
            return await FileExtractor.extract_from_docx(file_content)
        else:
            raise ValueError(f"Unsupported file type: {content_type}")
